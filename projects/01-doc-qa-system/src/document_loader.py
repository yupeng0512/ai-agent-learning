"""
文档加载器模块

支持的文档格式：
- PDF
- Markdown
- TXT
- DOCX
"""

import os
from typing import List, Optional
from pathlib import Path

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter


class DocumentLoader:
    """统一的文档加载器"""
    
    def __init__(
        self,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
    ):
        """
        初始化文档加载器
        
        Args:
            chunk_size: 文档切分块大小
            chunk_overlap: 块之间的重叠大小
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", "。", "！", "？", ".", "!", "?", " ", ""],
        )
        
        # 支持的文件格式
        self._supported_formats = [".pdf", ".md", ".txt", ".docx"]
    
    def get_supported_formats(self) -> List[str]:
        """获取支持的文件格式列表"""
        return self._supported_formats.copy()
    
    def load_file(self, file_path: str) -> List[Document]:
        """
        加载单个文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            文档列表
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        suffix = path.suffix.lower()
        
        if suffix == ".pdf":
            return self._load_pdf(file_path)
        elif suffix == ".md":
            return self._load_markdown(file_path)
        elif suffix == ".txt":
            return self._load_text(file_path)
        elif suffix == ".docx":
            return self._load_docx(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {suffix}")
    
    def load_directory(self, dir_path: str, extensions: Optional[List[str]] = None) -> List[Document]:
        """
        加载目录下的所有文档
        
        Args:
            dir_path: 目录路径
            extensions: 要加载的文件扩展名列表，默认全部支持的格式
            
        Returns:
            文档列表
        """
        if extensions is None:
            extensions = [".pdf", ".md", ".txt", ".docx"]
        
        documents = []
        dir_path = Path(dir_path)
        
        for ext in extensions:
            for file_path in dir_path.glob(f"**/*{ext}"):
                try:
                    docs = self.load_file(str(file_path))
                    documents.extend(docs)
                    print(f"✅ 已加载: {file_path.name} ({len(docs)} 块)")
                except Exception as e:
                    print(f"❌ 加载失败: {file_path.name} - {e}")
        
        return documents
    
    def load_text_content(self, content: str, metadata: Optional[dict] = None) -> List[Document]:
        """
        直接加载文本内容
        
        Args:
            content: 文本内容
            metadata: 元数据
            
        Returns:
            文档列表
        """
        if metadata is None:
            metadata = {"source": "direct_input"}
        
        doc = Document(page_content=content, metadata=metadata)
        return self.text_splitter.split_documents([doc])
    
    def _load_pdf(self, file_path: str) -> List[Document]:
        """加载 PDF 文件"""
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("请安装 pypdf: pip install pypdf")
        
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        
        doc = Document(
            page_content=text,
            metadata={"source": file_path, "type": "pdf"}
        )
        return self.text_splitter.split_documents([doc])
    
    def _load_markdown(self, file_path: str) -> List[Document]:
        """加载 Markdown 文件"""
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
        
        doc = Document(
            page_content=content,
            metadata={"source": file_path, "type": "markdown"}
        )
        return self.text_splitter.split_documents([doc])
    
    def _load_text(self, file_path: str) -> List[Document]:
        """加载纯文本文件"""
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
        
        doc = Document(
            page_content=content,
            metadata={"source": file_path, "type": "text"}
        )
        return self.text_splitter.split_documents([doc])
    
    def _load_docx(self, file_path: str) -> List[Document]:
        """加载 Word 文档"""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")
        
        doc = DocxDocument(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        
        document = Document(
            page_content=text,
            metadata={"source": file_path, "type": "docx"}
        )
        return self.text_splitter.split_documents([document])


# ============================================================
# 测试代码
# ============================================================

if __name__ == "__main__":
    loader = DocumentLoader(chunk_size=200, chunk_overlap=20)
    
    # 测试直接加载文本
    test_content = """
    # AI Agent 简介
    
    AI Agent 是一种能够自主感知环境、做出决策并采取行动的智能系统。
    
    ## 核心组件
    
    1. LLM（大语言模型）：作为"大脑"进行推理
    2. 工具：与外部世界交互的能力
    3. 记忆：存储历史信息
    
    ## 应用场景
    
    - 智能客服
    - 自动化办公
    - 代码助手
    """
    
    docs = loader.load_text_content(test_content, {"source": "test"})
    print(f"\n切分结果: {len(docs)} 块")
    for i, doc in enumerate(docs):
        print(f"\n--- Chunk {i+1} ---")
        print(doc.page_content[:100] + "...")
